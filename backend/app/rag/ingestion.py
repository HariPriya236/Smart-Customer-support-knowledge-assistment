import os
import pandas as pd
from typing import List, Dict, Any
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.core.config import settings

class DocumentIngestor:
    """
    Parses files (PDF, DOCX, TXT, MD, CSV) into standardized LangChain Document objects
    and splits them using RecursiveCharacterTextSplitter.
    """
    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )

    def load_and_split(self, file_path: str, filename: str, doc_category: str = "General Knowledge", doc_id: str = "") -> List[Document]:
        ext = os.path.splitext(filename)[1].lower()
        raw_documents = []

        if ext == ".pdf":
            raw_documents = self._parse_pdf(file_path, filename)
        elif ext in [".docx", ".doc"]:
            raw_documents = self._parse_docx(file_path, filename)
        elif ext in [".txt", ".md"]:
            raw_documents = self._parse_text(file_path, filename)
        elif ext == ".csv":
            raw_documents = self._parse_csv(file_path, filename)
        else:
            raw_documents = self._parse_text(file_path, filename)

        # Attach standard metadata
        chunks = self.splitter.split_documents(raw_documents)
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "doc_id": doc_id,
                "filename": filename,
                "file_type": ext.lstrip("."),
                "doc_category": doc_category,
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
        
        return chunks

    def _parse_pdf(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"page": page_num + 1, "source": filename}))
        except Exception:
            # Fallback to plain text read if PDF reader fails
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                docs.append(Document(page_content=f.read(), metadata={"source": filename, "page": 1}))
        return docs

    def _parse_docx(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            text = "\n".join(full_text)
            docs.append(Document(page_content=text, metadata={"source": filename, "page": 1}))
        except Exception:
            docs.append(Document(page_content=f"Content from {filename}", metadata={"source": filename, "page": 1}))
        return docs

    def _parse_text(self, file_path: str, filename: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [Document(page_content=content, metadata={"source": filename, "page": 1})]

    def _parse_csv(self, file_path: str, filename: str) -> List[Document]:
        docs = []
        try:
            df = pd.read_csv(file_path)
            # Format each row into key-value pairs
            row_texts = []
            for index, row in df.iterrows():
                row_str = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                row_texts.append(f"Row {index + 1}: {row_str}")
            full_text = "\n".join(row_texts)
            docs.append(Document(page_content=full_text, metadata={"source": filename, "page": 1}))
        except Exception:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                docs.append(Document(page_content=f.read(), metadata={"source": filename, "page": 1}))
        return docs
